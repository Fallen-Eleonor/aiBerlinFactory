from __future__ import annotations

from pathlib import Path

from docx import Document

from app.models import AnalyzeRequest, CapTableOutput, LegalOutput


class DocumentService:
    def __init__(self, base_dir: Path) -> None:
        self.base_dir = base_dir

    def generate_all(self, job_id: str, request: AnalyzeRequest, legal: LegalOutput, cap_table: CapTableOutput) -> dict[str, Path]:
        output_dir = self.base_dir / job_id
        output_dir.mkdir(parents=True, exist_ok=True)

        return {
            "gesellschaftsvertrag": self.generate_gesellschaftsvertrag(output_dir, request, legal, cap_table),
            "founder-resolution-summary": self.generate_founder_resolution_summary(output_dir, request, legal, cap_table),
            "handelsregister-checklist": self.generate_handelsregister_checklist(output_dir, request, legal),
        }

    def generate_gesellschaftsvertrag(self, output_dir: Path, request: AnalyzeRequest, legal: LegalOutput, cap_table: CapTableOutput) -> Path:
        file_path = output_dir / "gesellschaftsvertrag.docx"

        document = Document()
        document.add_heading("Gesellschaftsvertrag", level=1)
        document.add_paragraph(f"Firma: {request.company_name}")
        document.add_paragraph(f"Rechtsform-Empfehlung: {legal.recommended_entity}")
        document.add_paragraph(f"Stammkapital / geplantes Startkapital: EUR {request.available_capital_eur:,.0f}".replace(",", "."))
        document.add_paragraph(f"Branche: {request.industry}")
        document.add_paragraph(f"Bundesland: {request.bundesland}")
        document.add_paragraph("")
        document.add_paragraph("Muster - nicht rechtsverbindlich ohne Notar")
        document.add_paragraph(
            "Dieses Dokument ist ein hackathon-tauglicher Entwurf fuer die notarielle Vorbereitung und ersetzt keine individuelle Rechtsberatung."
        )
        document.add_paragraph("")
        document.add_paragraph("Empfohlene Gruendungslogik")
        document.add_paragraph(legal.reasoning)
        for bullet in legal.entity_rationale:
            document.add_paragraph(bullet, style="List Bullet")
        document.add_paragraph("")
        document.add_paragraph("Gruendungsschritte")
        for index, step in enumerate(legal.incorporation_steps, start=1):
            line = f"{index}. {step.title} - {step.description} ({step.owner}, {step.eta})"
            document.add_paragraph(line, style="List Number")
        document.add_paragraph("")
        document.add_paragraph("Nach der Eintragung")
        for item in legal.post_incorporation_checklist:
            document.add_paragraph(f"{item.title} - {item.description}", style="List Bullet")
        document.add_paragraph("")
        document.add_paragraph("Hinweis zur weiteren Struktur")
        document.add_paragraph(legal.conversion_note)
        document.add_paragraph("")
        document.add_paragraph("Vorgeschlagene Eigentumsstruktur")
        for allocation in cap_table.allocations:
            document.add_paragraph(
                f"{allocation.holder}: {allocation.ownership_percent:.1f}% ({allocation.role})",
                style="List Bullet",
            )
        document.add_paragraph(f"ESOP / Mitarbeiterpool: {cap_table.option_pool_percent:.1f}%")
        document.add_paragraph(f"Beraterpool: {cap_table.advisor_pool_percent:.1f}%")
        document.add_paragraph("Geschaeftsfuehrer (Platzhalter): Founder 1, vorbehaltlich finaler Gesellschafterbeschluesse.")

        document.save(file_path)
        return file_path

    def generate_founder_resolution_summary(self, output_dir: Path, request: AnalyzeRequest, legal: LegalOutput, cap_table: CapTableOutput) -> Path:
        file_path = output_dir / "founder-resolution-summary.txt"
        content = "\n".join(
            [
                f"Founder Resolution Summary: {request.company_name}",
                "",
                f"Recommended entity: {legal.recommended_entity}",
                f"Industry: {request.industry}",
                f"Bundesland: {request.bundesland}",
                f"Founder count: {request.founder_count}",
                f"Capital available: EUR {request.available_capital_eur:,.0f}".replace(",", "."),
                "",
                "Why this route:",
                *[f"- {item}" for item in legal.entity_rationale],
                "",
                "Immediate board-level decisions:",
                "- Confirm final shareholder split before the notary draft.",
                "- Confirm managing director appointment wording.",
                "- Confirm the target bank and capital deposit timing.",
                "",
                "Draft ownership split:",
                *[f"- {allocation.holder}: {allocation.ownership_percent:.1f}% ({allocation.role})" for allocation in cap_table.allocations],
                f"- ESOP reserve: {cap_table.option_pool_percent:.1f}%",
                f"- Advisor reserve: {cap_table.advisor_pool_percent:.1f}%",
            ]
        )
        file_path.write_text(content, encoding="utf-8")
        return file_path

    def generate_handelsregister_checklist(self, output_dir: Path, request: AnalyzeRequest, legal: LegalOutput) -> Path:
        file_path = output_dir / "handelsregister-checklist.txt"
        lines = [
            f"Handelsregister Checklist: {request.company_name}",
            "",
            "Preparation items:",
        ]
        lines.extend(f"- {step.title}: {step.description}" for step in legal.incorporation_steps)
        lines.extend(
            [
                "",
                "Post-registration items:",
            ]
        )
        lines.extend(f"- {item.title}: {item.description}" for item in legal.post_incorporation_checklist)
        file_path.write_text("\n".join(lines), encoding="utf-8")
        return file_path
